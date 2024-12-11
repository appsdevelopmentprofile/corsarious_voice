import time
import streamlit as st
import os
import subprocess
from pydub import AudioSegment
import speech_recognition as sr
import io
from docx import Document

# Set the FFMPEG_BINARY to the direct path of ffmpeg executable
os.environ["FFMPEG_BINARY"] = "/usr/local/bin/ffmpeg"
os.environ["FFPROBE_BINARY"] = "/usr/local/bin/ffprobe"

# Function to check if ffmpeg is accessible
def check_ffmpeg():
    try:
        # Run ffmpeg command to check its version
        result = subprocess.run(['sudo', 'ffmpeg', '-version'], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)

        # Display the output in Streamlit
        ffmpeg_version = result.stdout
        st.write("FFmpeg version detected successfully:")
        st.write(ffmpeg_version)
    except FileNotFoundError:
        st.error("FFmpeg is not installed or not found in the specified path.")
    except PermissionError:
        st.error("Permission denied while trying to run ffmpeg with sudo.")
    except Exception as e:
        st.error(f"An error occurred: {e}")

# Streamlit app title
st.title("FFmpeg and Speech-to-Text Integration Example")

# Button to check FFmpeg version
if st.button("Check FFmpeg Version"):
    check_ffmpeg()

# Function 1: Play "electric_unit_heater.wav" file from GitHub repo (local directory)
def play_electric_unit_heater():
    st.header("Function 1: Play 'electric_unit_heater.wav'")
    file_path = "electric_unit_heater.wav"
    if os.path.exists(file_path):
        st.audio(file_path, format="audio/wav")
    else:
        st.error("File 'electric_unit_heater.wav' not found!")

# Function 2: Record voice, save as wav, and allow playback
def record_voice():
    st.header("Function 2: Record Voice")

    # Upload an audio file (only wav or mp3)
    uploaded_file = st.file_uploader("Choose a file", type=["wav", "mp3"])
       
    if uploaded_file is not None:
        # Convert uploaded audio to AudioSegment
        audio = AudioSegment.from_file(uploaded_file)
           
        # Save the uploaded audio as WAV
        wav_file_path = "uploaded_audio.wav"
        audio.export(wav_file_path, format="wav")
           
        # Display success message and play the audio
        st.success(f"Audio uploaded and saved as {wav_file_path}!")
        st.audio(wav_file_path, format="audio/wav")
           
        # Optionally save as MP3 as well
        mp3_file_path = "uploaded_audio.mp3"
        audio.export(mp3_file_path, format="mp3")
        st.audio(mp3_file_path, format="audio/mp3")

# Function 3: Convert speech from wav file to text using Google Speech Recognition
def recognize_speech_from_wav(wav_file):
    recognizer = sr.Recognizer()

    # Load the .wav file
    with sr.AudioFile(wav_file) as source:
        audio = recognizer.record(source)
        try:
            # Use Google Web Speech API to recognize the speech
            text = recognizer.recognize_google(audio)
            return text.lower()
        except sr.UnknownValueError:
            return "Speech not recognized"
        except sr.RequestError as e:
            return f"Error with the request: {e}"

# Function 4: Process the "electric_unit_heater.wav" file
def process_electric_unit_heater():
    st.header("Function 4: Recognize Speech from 'electric_unit_heater.wav'")

    wav_file = "electric_unit_heater.wav"
   
    if os.path.exists(wav_file):
        # Convert the speech in the .wav file to text
        recognized_text = recognize_speech_from_wav(wav_file)
        st.write(f"Recognized Text: {recognized_text}")
    else:
        st.error(f"File '{wav_file}' not found!")


# Function 4: Show progress bar for task completion
def show_progress_bar():
    st.header("Function 4: Progress Bar for Task Completion")

    progress_bar = st.progress(0)
    for i in range(1, 101):
        progress_bar.progress(i)
        time.sleep(0.1)  # Simulating a process with a delay
    st.success("Task Completed!")

# Function 5: Generate a checklist document after phases 1-4
import speech_recognition as sr
from docx import Document
from pydub import AudioSegment
import io

# Function to recognize speech from the WAV file
def recognize_speech_from_wav(file_path):
    recognizer = sr.Recognizer()
    
    # Export the audio as raw audio data in bytes (for speech_recognition to process)
    audio_data = io.BytesIO()
    audio = AudioSegment.from_wav(file_path)
    audio.export(audio_data, format="wav")
    audio_data.seek(0)  # Go to the beginning of the audio data

    # Recognize speech from the audio data
    with sr.AudioFile(audio_data) as source:
        audio_recorded = recognizer.record(source)
        try:
            text = recognizer.recognize_google(audio_recorded)
            return text.lower()  # Return the recognized text in lowercase
        except sr.UnknownValueError:
            return "unrecognized"
        except sr.RequestError as e:
            return f"error: {e}"

# Function to create a checklist document from recognized speech
def create_checklist_doc_from_speech(text):
    # Create a new Document
    doc = Document()

    # Title of the document
    doc.add_heading('Field Engineer Questionnaire', 0)

    # Add a table with 2 columns
    table = doc.add_table(rows=1, cols=2)

    # Set the table header (first row)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Question'
    hdr_cells[1].text = 'Answer (YES/NO)'

    # Split the recognized text into checklist items (basic split by periods for demo)
    lines = text.split(".")  # Split by periods for simplicity, customize as needed

    # Add questions and the YES/NO options in the table
    for line in lines:
        line = line.strip()
        if line:
            row_cells = table.add_row().cells
            row_cells[0].text = line  # Add the question
            row_cells[1].text = '☐ YES ☐ NO'  # Placeholder for the "YES/NO" selection

    # Save the document
    doc.save('field_engineer_checklist_from_speech.docx')
    print("Checklist document created successfully based on speech recognition.")

# Main function to process the WAV file and create the document
def process_audio_and_create_doc(file_path):
    print("Recognizing speech from audio file...")
    recognized_text = recognize_speech_from_wav(file_path)

    print("Creating checklist document...")
    create_checklist_doc_from_speech(recognized_text)

# Path to the WAV file (engineer_equipment.wav)
file_path = "engineer_equipment.wav"

# Run the process
process_audio_and_create_doc(file_path)


# Main App
st.title("Audio Demo App with Speech-to-Text")

# Sequential execution of functions
if st.button("Start Function 1: Play 'electric_unit_heater.wav'"):
    play_electric_unit_heater()

if st.button("Start Function 2: Record Voice"):
    record_voice()

if st.button("Start Function 3: Recognize Speech from 'electric_unit_heater.wav'"):
    process_electric_unit_heater()
    
if st.button("Start Function 4: Progress Bar"):
    show_progress_bar()

if st.button("Start Function 5: Generate Checklist Document"):
    generate_checklist_doc()
